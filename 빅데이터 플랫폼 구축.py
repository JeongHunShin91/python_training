# 1) 첫 번째 문단
# “파이썬을 활용한 Word문서 작성” 텍스트
# Font: 맑은 고딕
# Font size: 24
# Bold

from docx.shared import Pt
from docx.oxml.ns import qn
from docx import Document
filename = '신정훈.docx'
document = Document()

document = Document(filename)
document.add_paragraph('파이썬을 활용한 Word문서 작성')

first_paragraph = document.paragraphs[0]
first_run = first_paragraph.runs[0]

first_run.bold = True
first_run.font.size = Pt(24)

first_run.font.name = '맑은 고딕'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

document.save(filename)

# [문항2]  2) 두 번째 문단
# “본인이름” 텍스트
# Font: 굴림
# Font size: 20
# Underline

from docx.shared import Pt
from docx.oxml.ns import qn

document = Document(filename)
document.add_paragraph('신정훈')
second_paragraph  = document.paragraphs[1]
second_run = second_paragraph.runs[0]

second_run.underline = True
second_run.font.size = Pt(20)

second_run.font.name = '굴림'
second_run._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')

document.save(filename)
# [문항3]  3) 세 번째 문단
# Tjoeun_logo2.jpg 이미지 파일 첨부
# 중앙 정렬
# 그림 설명 부분 [더조은 컴퓨터 아카데미] 추가
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK

document = Document(filename)

new_paragraph = document.add_paragraph()

new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

logo_run = new_paragraph.add_run()

logo_run.add_picture('tjoeun_logo2.jpg')

logo_run.add_break(WD_BREAK.LINE)

caption_run = new_paragraph.add_run('[더조은 컴퓨터 아카데미]')

document.save(filename)
# [문항4]  4) 네 번째 문단
# “훈련중입니다.” 텍스트
# Font: 맑은고딕
# Font size: 40
# 이탤릭체

document = Document(filename)
document.add_paragraph('훈련중입니다.')
next_paragraph  = document.paragraphs[3]
next_run = next_paragraph.runs[0]

next_run.italic = True
next_run.font.size = Pt(40)

next_run.font.name = '맑은고딕'
next_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은고딕')

document.save(filename)